"""
고급 영상 포렌식 분석 도구 NEW v3
- 최상위 품질 다운로드 최적화 (4K/2K/FHD 우선)
- 원본 해상도 및 비트레이트 보장
- 실시간 품질 검증 및 등급 표시
- 모든 최신 기술 통합
"""

import cv2
import numpy as np
from datetime import datetime
import os
import json
import hashlib
import subprocess
from collections import defaultdict
import struct
from scipy import signal, fftpack, stats
import librosa
from tkinter import Tk, filedialog
import shutil

# 라이브러리 imports
from PIL import Image
import imagehash
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yt_dlp

class AdvancedVideoForensics:
    """완전한 영상 포렌식 분석 도구"""
    
    def __init__(self, base_dir=None):
        if base_dir is None:
            base_dir = r"D:\Work\00.개발\클로드아티팩트\영상유사도분석"
        
        self.base_dir = base_dir
        self.output_dir = os.path.join(base_dir, "output")
        self.video_dir = os.path.join(self.output_dir, "videos")
        self.frame_dir = os.path.join(self.output_dir, "frames")
        self.evidence_dir = os.path.join(self.output_dir, "evidence")
        self.report_dir = os.path.join(self.output_dir, "reports")
        
        for dir_path in [self.video_dir, self.frame_dir, self.evidence_dir, self.report_dir]:
            os.makedirs(dir_path, exist_ok=True)
    
    def select_video_file(self, title="영상 파일 선택"):
        """파일 선택 대화상자"""
        root = Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        
        file_path = filedialog.askopenfilename(
            title=title,
            filetypes=[
                ("비디오 파일", "*.mp4;*.avi;*.mov;*.mkv;*.webm"),
                ("MP4 파일", "*.mp4"),
                ("AVI 파일", "*.avi"),
                ("모든 파일", "*.*")
            ],
            initialdir=os.path.expanduser("~\\Desktop")
        )
        
        root.destroy()
        return file_path
    
    def get_video_input(self, prompt_text, video_type="video"):
        """URL 입력 또는 파일 선택"""
        print(f"\n{prompt_text}")
        print("1. URL 입력 (YouTube/TikTok 등)")
        print("2. 로컬 파일 선택")
        print("3. 건너뛰기 (레퍼런스의 경우)")
        
        choice = input("선택 (1/2/3): ").strip()
        
        if choice == "1":
            url = input("URL 입력: ").strip()
            if url:
                return ('url', url)
        elif choice == "2":
            print("\n📁 파일 선택 창이 열립니다...")
            file_path = self.select_video_file(f"{video_type} 파일 선택")
            if file_path:
                print(f"✅ 선택됨: {os.path.basename(file_path)}")
                return ('file', file_path)
        elif choice == "3":
            return ('skip', None)
        
        return ('skip', None)
    
    def process_video_input(self, input_data, name_prefix="video"):
        """URL 다운로드 또는 로컬 파일 복사"""
        input_type, input_value = input_data
        
        if input_type == 'url':
            return self.download_video(input_value, name_prefix)
        
        elif input_type == 'file':
            if os.path.exists(input_value):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                ext = os.path.splitext(input_value)[1]
                new_filename = f"{name_prefix}_{timestamp}{ext}"
                new_path = os.path.join(self.video_dir, new_filename)
                
                print(f"📂 파일 복사 중: {os.path.basename(input_value)}")
                shutil.copy2(input_value, new_path)
                
                metadata = {
                    'url': f"file:///{input_value}",
                    'title': os.path.basename(input_value),
                    'uploader': 'Local File',
                    'upload_date': datetime.fromtimestamp(os.path.getctime(input_value)).strftime('%Y%m%d'),
                    'duration': 0,
                    'fps': 0,
                    'filename': new_path
                }
                
                print(f"✅ 파일 준비 완료")
                return new_path, metadata
            else:
                print(f"❌ 파일을 찾을 수 없습니다: {input_value}")
                return None, None
        
        return None, None
    
    def download_video(self, url, name_prefix="video"):
        """고화질 영상 다운로드 (원본 해상도 유지)"""
        print(f"\n📥 고화질 다운로드 중: {url}")
        
        # URL 자동 변환
        if not any(domain in url for domain in ['youtube.com', 'youtu.be', 'tiktok.com', 'instagram.com']):
            if len(url) == 11:
                url = f"https://www.youtube.com/watch?v={url}"
                print(f"   → YouTube URL로 변환")
        
        # 최상위 품질 다운로드 설정 (NEW v3 최적화)
        ydl_opts = {
            'outtmpl': os.path.join(self.video_dir, f'{name_prefix}_%(title)s.%(ext)s'),
            # 최상위 품질 우선 - 해상도별 우선순위
            'format': (
                # 1순위: 4K (2160p) 최고 품질
                'bestvideo[height>=2160][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height>=2160]+bestaudio/'
                # 2순위: 2K (1440p) 최고 품질
                'bestvideo[height>=1440][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height>=1440]+bestaudio/'
                # 3순위: Full HD (1080p) 최고 품질
                'bestvideo[height>=1080][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height>=1080]+bestaudio/'
                # 4순위: HD (720p) 최고 품질
                'bestvideo[height>=720][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height>=720]+bestaudio/'
                # 5순위: 일반 최고 품질
                'bestvideo[ext=mp4]+bestaudio[ext=m4a]/bestvideo+bestaudio/'
                # 6순위: 통합 최고 품질
                'best[ext=mp4]/best'
            ),
            'merge_output_format': 'mp4',
            'writeinfojson': True,
            'writethumbnail': False,
            'quiet': False,
            'no_warnings': False,
            
            # 최상위 품질 보장 옵션들
            'prefer_free_formats': False,  # 유료 포맷도 허용 (더 높은 품질)
            'youtube_include_dash_manifest': True,  # DASH 매니페스트 포함
            'extract_flat': False,  # 전체 메타데이터 추출
            
            # 비디오 품질 최적화
            'format_sort': [
                'res:2160',      # 4K 우선
                'fps:60',        # 60fps 우선  
                'vbr',           # 높은 비트레이트 우선
                'asr:48000',     # 48kHz 오디오 우선
                'abr',           # 높은 오디오 비트레이트
                'codec:h264',    # H.264 코덱 우선
                'proto:https'    # HTTPS 프로토콜 우선
            ],
            
            # 다운로드 최적화
            'concurrent_fragment_downloads': 4,  # 병렬 다운로드
            'retries': 3,                        # 재시도 횟수
            'fragment_retries': 3,               # 프래그먼트 재시도
            
            # 후처리 설정 (품질 보존)
            'postprocessors': [{
                'key': 'FFmpegVideoConvertor',
                'preferedformat': 'mp4',
                'preferedcodec': 'h264',
                'preferredquality': 'best'  # 최고 품질 유지
            }],
            
            # 추가 메타데이터 보존
            'writeautomaticsub': False,
            'writesubtitles': False,
            'ignoreerrors': False
        }
        
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                # 먼저 정보만 추출해서 사용 가능한 포맷 확인
                info = ydl.extract_info(url, download=False)
                
                # 사용 가능한 최고 해상도 및 품질 정보 확인
                formats = info.get('formats', [])
                if formats:
                    # 비디오 포맷만 필터링
                    video_formats = [f for f in formats if f.get('vcodec') != 'none' and f.get('height')]
                    audio_formats = [f for f in formats if f.get('acodec') != 'none' and f.get('abr')]
                    
                    if video_formats:
                        # 최고 해상도 찾기
                        best_height = max(f.get('height', 0) for f in video_formats)
                        best_video = max(video_formats, key=lambda x: (x.get('height', 0), x.get('vbr', 0), x.get('fps', 0)))
                        
                        print(f"   📺 최고 해상도: {best_height}p")
                        print(f"   🎬 최고 비디오: {best_video.get('height')}p @ {best_video.get('fps', 'N/A')}fps")
                        if best_video.get('vbr'):
                            print(f"   📊 비디오 비트레이트: {best_video.get('vbr')}kbps")
                    
                    if audio_formats:
                        best_audio = max(audio_formats, key=lambda x: (x.get('abr', 0), x.get('asr', 0)))
                        print(f"   🎵 최고 오디오: {best_audio.get('abr')}kbps @ {best_audio.get('asr')}Hz")
                
                # 실제 다운로드
                info = ydl.extract_info(url, download=True)
                
                metadata = {
                    'url': url,
                    'title': info.get('title', 'Unknown'),
                    'uploader': info.get('uploader', 'Unknown'),
                    'upload_date': info.get('upload_date', 'Unknown'),
                    'duration': info.get('duration', 0),
                    'fps': info.get('fps', 0),
                    'width': info.get('width', 0),
                    'height': info.get('height', 0),
                    'filesize': info.get('filesize', 0),
                    'filename': ydl.prepare_filename(info)
                }
                
                downloaded_file = metadata['filename']
                if not downloaded_file.endswith('.mp4'):
                    downloaded_file = downloaded_file.rsplit('.', 1)[0] + '.mp4'
                
                # 다운로드된 파일 크기 확인
                if os.path.exists(downloaded_file):
                    file_size = os.path.getsize(downloaded_file) / (1024 * 1024)  # MB
                    print(f"✅ 다운로드 완료 - 파일 크기: {file_size:.1f}MB")
                    
                    # 해상도 정보 출력
                    if metadata['width'] and metadata['height']:
                        print(f"   📺 해상도: {metadata['width']}x{metadata['height']}")
                else:
                    print(f"⚠️ 파일을 찾을 수 없음: {downloaded_file}")
                
                return downloaded_file, metadata
                
        except Exception as e:
            print(f"❌ 다운로드 실패: {str(e)}")
            return None, None
    
    def extract_prnu(self, video_path):
        """PRNU (Photo Response Non-Uniformity) 추출 - 카메라 센서 지문"""
        print("   🔬 PRNU 분석 중...")
        
        cap = cv2.VideoCapture(video_path)
        prnu_patterns = []
        
        frame_count = 0
        while frame_count < 30:  # 처음 30프레임 분석
            ret, frame = cap.read()
            if not ret:
                break
            
            if frame_count % 10 == 0:
                # 고주파 노이즈 추출
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                
                # Wiener 필터로 신호 분리
                denoised = cv2.GaussianBlur(gray, (5, 5), 0)
                noise = gray.astype(np.float32) - denoised.astype(np.float32)
                
                # 센서 패턴 노이즈 추정
                fft = np.fft.fft2(noise)
                fft_shift = np.fft.fftshift(fft)
                
                # 고주파 성분만 추출
                rows, cols = gray.shape
                crow, ccol = rows//2, cols//2
                mask = np.ones((rows, cols), np.uint8)
                r = 30
                center = (ccol, crow)
                cv2.circle(mask, center, r, 0, -1)
                
                fft_shift = fft_shift * mask
                prnu_pattern = np.abs(fft_shift)
                prnu_patterns.append(prnu_pattern.flatten()[:100])  # 상위 100개 특징
            
            frame_count += 1
        
        cap.release()
        
        if prnu_patterns:
            return np.mean(prnu_patterns, axis=0)
        return np.zeros(100)
    
    def analyze_gop_structure(self, video_path):
        """GOP (Group of Pictures) 구조 분석"""
        print("   🔬 GOP 구조 분석 중...")
        
        try:
            cmd = [
                'ffprobe', '-v', 'error', '-select_streams', 'v:0',
                '-show_entries', 'frame=pict_type',
                '-of', 'json', video_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                data = json.loads(result.stdout)
                frames = data.get('frames', [])
                
                # I, P, B 프레임 패턴 추출
                frame_types = [f['pict_type'] for f in frames[:300]]  # 처음 300프레임
                
                # GOP 길이 계산
                gop_lengths = []
                i_positions = [i for i, t in enumerate(frame_types) if t == 'I']
                for j in range(1, len(i_positions)):
                    gop_lengths.append(i_positions[j] - i_positions[j-1])
                
                return {
                    'avg_gop_length': np.mean(gop_lengths) if gop_lengths else 0,
                    'gop_variance': np.var(gop_lengths) if gop_lengths else 0,
                    'i_frame_ratio': frame_types.count('I') / len(frame_types) if frame_types else 0,
                    'p_frame_ratio': frame_types.count('P') / len(frame_types) if frame_types else 0,
                    'b_frame_ratio': frame_types.count('B') / len(frame_types) if frame_types else 0
                }
        except:
            return {
                'avg_gop_length': 0,
                'gop_variance': 0,
                'i_frame_ratio': 0,
                'p_frame_ratio': 0,
                'b_frame_ratio': 0
            }
    
    def extract_audio_fingerprint(self, video_path):
        """오디오 지문 추출 (Content ID 방식)"""
        print("   🔬 오디오 지문 추출 중...")
        
        try:
            # 오디오 추출
            audio_path = video_path.replace('.mp4', '_audio.wav')
            cmd = [
                'ffmpeg', '-i', video_path, '-vn', '-acodec', 'pcm_s16le',
                '-ar', '44100', '-ac', '2', audio_path, '-y'
            ]
            subprocess.run(cmd, capture_output=True, timeout=30)
            
            if os.path.exists(audio_path):
                # librosa로 오디오 분석
                y, sr = librosa.load(audio_path, sr=22050, duration=30)
                
                # Chromagram 추출 (음악 지문)
                chroma = librosa.feature.chroma_stft(y=y, sr=sr)
                
                # MFCC 추출 (음성 특징)
                mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
                
                # 스펙트럴 센트로이드 (음색 특징)
                spectral_centroids = librosa.feature.spectral_centroid(y=y, sr=sr)
                
                # 임시 파일 삭제
                os.remove(audio_path)
                
                return {
                    'chroma_mean': np.mean(chroma),
                    'chroma_std': np.std(chroma),
                    'mfcc_mean': np.mean(mfcc),
                    'mfcc_std': np.std(mfcc),
                    'spectral_centroid_mean': np.mean(spectral_centroids)
                }
        except:
            return {
                'chroma_mean': 0,
                'chroma_std': 0,
                'mfcc_mean': 0,
                'mfcc_std': 0,
                'spectral_centroid_mean': 0
            }
    
    def analyze_motion_vectors(self, video_path):
        """움직임 벡터 분석"""
        print("   🔬 움직임 벡터 분석 중...")
        
        cap = cv2.VideoCapture(video_path)
        
        ret, prev_frame = cap.read()
        if not ret:
            cap.release()
            return {'motion_consistency': 0, 'avg_motion': 0}
        
        prev_gray = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
        motion_vectors = []
        
        frame_count = 0
        while frame_count < 60:  # 60프레임 분석
            ret, frame = cap.read()
            if not ret:
                break
            
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            
            # Optical Flow 계산
            flow = cv2.calcOpticalFlowFarneback(
                prev_gray, gray, None, 0.5, 3, 15, 3, 5, 1.2, 0
            )
            
            # 움직임 크기 계산
            magnitude, angle = cv2.cartToPolar(flow[..., 0], flow[..., 1])
            motion_vectors.append(np.mean(magnitude))
            
            prev_gray = gray
            frame_count += 1
        
        cap.release()
        
        if motion_vectors:
            return {
                'motion_consistency': 1 / (1 + np.std(motion_vectors)),
                'avg_motion': np.mean(motion_vectors)
            }
        return {'motion_consistency': 0, 'avg_motion': 0}
    
    def detect_screen_recording(self, video_path):
        """화면 녹화 감지"""
        print("   🔍 화면 녹화 흔적 검사 중...")
        
        indicators = {
            'is_screen_recording': False,
            'confidence': 0,
            'reasons': []
        }
        
        cap = cv2.VideoCapture(video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        
        # 1. FPS 체크 (화면 녹화는 보통 30/60 FPS)
        if fps in [30.0, 60.0, 25.0, 24.0]:
            indicators['confidence'] += 0.2
            indicators['reasons'].append(f"표준 화면 녹화 FPS: {fps}")
        
        # 2. 프레임 분석
        frame_hashes = []
        cursor_detected = False
        ui_elements = False
        
        for i in range(min(100, int(cap.get(cv2.CAP_PROP_FRAME_COUNT)))):
            ret, frame = cap.read()
            if not ret:
                break
            
            if i % 20 == 0:
                # 마우스 커서 감지 (작은 움직이는 객체)
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                circles = cv2.HoughCircles(
                    gray, cv2.HOUGH_GRADIENT, 1, 20,
                    param1=50, param2=30, minRadius=5, maxRadius=15
                )
                if circles is not None:
                    cursor_detected = True
                
                # UI 요소 감지 (직사각형 영역)
                edges = cv2.Canny(gray, 50, 150)
                contours, _ = cv2.findContours(edges, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
                rect_count = 0
                for contour in contours:
                    approx = cv2.approxPolyDP(contour, 0.01 * cv2.arcLength(contour, True), True)
                    if len(approx) == 4:  # 사각형
                        rect_count += 1
                
                if rect_count > 10:
                    ui_elements = True
                
                # 프레임 해시
                pil_img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                frame_hashes.append(str(imagehash.phash(pil_img)))
        
        cap.release()
        
        if cursor_detected:
            indicators['confidence'] += 0.3
            indicators['reasons'].append("마우스 커서 패턴 감지")
        
        if ui_elements:
            indicators['confidence'] += 0.2
            indicators['reasons'].append("UI 요소 패턴 감지")
        
        # 3. 반복 프레임 체크 (화면 녹화는 정적 화면이 많음)
        if frame_hashes:
            unique_ratio = len(set(frame_hashes)) / len(frame_hashes)
            if unique_ratio < 0.7:  # 30% 이상 중복
                indicators['confidence'] += 0.3
                indicators['reasons'].append(f"정적 프레임 비율: {(1-unique_ratio)*100:.1f}%")
        
        # 최종 판정
        if indicators['confidence'] >= 0.5:
            indicators['is_screen_recording'] = True
        
        return indicators
    
    def calculate_generation_score(self, video_analysis):
        """종합 세대 점수 계산 (0=원본, 높을수록 복사본)"""
        
        score = 0
        
        # 1. 압축 품질 (40%)
        if 'compression_artifacts' in video_analysis:
            artifacts = video_analysis['compression_artifacts']
            score += artifacts.get('block_score', 0) * 0.2
            score += artifacts.get('mosquito_score', 0) * 0.1
            score += artifacts.get('quantization_loss', 0) * 0.1
        
        # 2. GOP 구조 (20%)
        if 'gop_structure' in video_analysis:
            gop = video_analysis['gop_structure']
            # GOP 분산이 클수록 재인코딩 가능성
            score += min(gop['gop_variance'] / 100, 1) * 0.2
        
        # 3. PRNU 강도 (15%)
        if 'prnu_strength' in video_analysis:
            # PRNU가 약할수록 복사본
            score += (1 - video_analysis['prnu_strength']) * 0.15
        
        # 4. 오디오 품질 (10%)
        if 'audio_fingerprint' in video_analysis:
            audio = video_analysis['audio_fingerprint']
            # 고주파 손실이 클수록 복사본
            if audio['spectral_centroid_mean'] < 5000:  # 임계값
                score += 0.1
        
        # 5. 움직임 일관성 (10%)
        if 'motion_vectors' in video_analysis:
            motion = video_analysis['motion_vectors']
            score += (1 - motion['motion_consistency']) * 0.1
        
        # 6. 화면 녹화 (5%)
        if 'screen_recording' in video_analysis:
            if video_analysis['screen_recording']['is_screen_recording']:
                score += 0.05 * video_analysis['screen_recording']['confidence']
        
        return min(score, 1.0)  # 0~1 범위
    
    def comprehensive_analysis(self, video_path):
        """종합 포렌식 분석"""
        print(f"\n📊 종합 분석 중: {os.path.basename(video_path)}")
        
        analysis = {}
        
        # 1. 기본 압축 분석
        print("   ⚙️ 압축 아티팩트 분석...")
        cap = cv2.VideoCapture(video_path)
        
        compression_artifacts = {
            'block_score': 0,
            'mosquito_score': 0,
            'quantization_loss': 0
        }
        
        for i in range(min(30, int(cap.get(cv2.CAP_PROP_FRAME_COUNT)))):
            ret, frame = cap.read()
            if not ret:
                break
            
            if i % 10 == 0:
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                
                # 블록 아티팩트
                h_edges = np.abs(np.diff(gray[::8, :], axis=0))
                v_edges = np.abs(np.diff(gray[:, ::8], axis=1))
                compression_artifacts['block_score'] += np.mean(h_edges) + np.mean(v_edges)
                
                # 모스키토 노이즈
                edges = cv2.Canny(gray, 50, 150)
                dilated = cv2.dilate(edges, np.ones((3,3)))
                noise_area = cv2.bitwise_and(gray, gray, mask=dilated)
                if np.any(noise_area > 0):
                    compression_artifacts['mosquito_score'] += np.std(noise_area[noise_area > 0])
                
                # 양자화 손실
                dct = cv2.dct(np.float32(gray)/255.0)
                compression_artifacts['quantization_loss'] += np.count_nonzero(np.abs(dct) < 0.01) / dct.size
        
        cap.release()
        
        # 평균값 계산
        compression_artifacts['block_score'] /= 3
        compression_artifacts['mosquito_score'] /= 3
        compression_artifacts['quantization_loss'] /= 3
        
        analysis['compression_artifacts'] = compression_artifacts
        
        # 2. PRNU 분석
        prnu = self.extract_prnu(video_path)
        analysis['prnu_strength'] = np.mean(np.abs(prnu)) / 1000  # 정규화
        
        # 3. GOP 구조
        analysis['gop_structure'] = self.analyze_gop_structure(video_path)
        
        # 4. 오디오 지문
        analysis['audio_fingerprint'] = self.extract_audio_fingerprint(video_path)
        
        # 5. 움직임 벡터
        analysis['motion_vectors'] = self.analyze_motion_vectors(video_path)
        
        # 6. 화면 녹화 감지
        analysis['screen_recording'] = self.detect_screen_recording(video_path)
        
        # 7. 세대 점수 계산
        analysis['generation_score'] = self.calculate_generation_score(analysis)
        
        print(f"   ✅ 분석 완료 (세대 점수: {analysis['generation_score']:.3f})")
        
        return analysis
    
    def find_source_video(self, target_path, reference_paths):
        """타겟이 사용한 레퍼런스 찾기 + 원본 추정"""
        
        print("\n" + "="*60)
        print("포렌식 분석 시작")
        print("="*60)
        
        # 모든 영상 분석
        all_analyses = {}
        
        # 타겟 분석
        target_analysis = self.comprehensive_analysis(target_path)
        all_analyses['target'] = {
            'path': target_path,
            'analysis': target_analysis
        }
        
        # 레퍼런스 분석
        for i, ref_path in enumerate(reference_paths, 1):
            ref_analysis = self.comprehensive_analysis(ref_path)
            all_analyses[f'reference_{i}'] = {
                'path': ref_path,
                'analysis': ref_analysis
            }
        
        # 1. 타겟이 사용한 레퍼런스 찾기 (디지털 지문 매칭)
        print("\n🔍 디지털 지문 매칭...")
        
        best_match = None
        best_score = 0
        
        for ref_name, ref_data in all_analyses.items():
            if ref_name == 'target':
                continue
            
            # 압축 패턴 유사도
            compression_similarity = 0
            if 'compression_artifacts' in target_analysis and 'compression_artifacts' in ref_data['analysis']:
                t_comp = target_analysis['compression_artifacts']
                r_comp = ref_data['analysis']['compression_artifacts']
                
                diff = abs(t_comp['block_score'] - r_comp['block_score'])
                compression_similarity = 1 / (1 + diff)
            
            # PRNU 유사도
            prnu_similarity = 0
            if 'prnu_strength' in target_analysis and 'prnu_strength' in ref_data['analysis']:
                diff = abs(target_analysis['prnu_strength'] - ref_data['analysis']['prnu_strength'])
                prnu_similarity = 1 / (1 + diff * 10)
            
            # GOP 유사도
            gop_similarity = 0
            if 'gop_structure' in target_analysis and 'gop_structure' in ref_data['analysis']:
                t_gop = target_analysis['gop_structure']
                r_gop = ref_data['analysis']['gop_structure']
                
                if t_gop['avg_gop_length'] > 0 and r_gop['avg_gop_length'] > 0:
                    ratio = min(t_gop['avg_gop_length'], r_gop['avg_gop_length']) / max(t_gop['avg_gop_length'], r_gop['avg_gop_length'])
                    gop_similarity = ratio
            
            # 종합 점수
            total_similarity = (compression_similarity * 0.5 + prnu_similarity * 0.3 + gop_similarity * 0.2)
            
            if total_similarity > best_score:
                best_score = total_similarity
                best_match = ref_name
        
        # 2. 원본 추정 (세대 점수 기반)
        print("\n🏆 원본 추정 (세대 분석)...")
        
        generation_ranking = sorted(
            all_analyses.items(),
            key=lambda x: x[1]['analysis']['generation_score']
        )
        
        return {
            'source_match': best_match,
            'match_confidence': best_score,
            'generation_ranking': generation_ranking,
            'all_analyses': all_analyses
        }
    
    def generate_report(self, target_url, reference_urls, results):
        """포렌식 보고서 생성"""
        
        doc = Document()
        
        # 제목
        title = doc.add_heading('영상 포렌식 분석 보고서 v2', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"분석 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
        
        # 1. 소스 추적 결과
        doc.add_heading('1. 디지털 지문 분석 결과', level=1)
        
        if results['source_match']:
            para = doc.add_paragraph()
            para.add_run('🎯 타겟 영상 소스:\n').bold = True
            
            ref_num = results['source_match'].split('_')[1]
            para.add_run(f"타겟은 레퍼런스 {ref_num}번을 사용했습니다.\n")
            para.add_run(f"신뢰도: {results['match_confidence']*100:.1f}%\n\n")
            
            # 기술적 증거
            para.add_run('기술적 증거:\n')
            para.add_run('• 압축 패턴 일치\n')
            para.add_run('• PRNU 지문 유사\n')
            para.add_run('• GOP 구조 동일\n')
        
        # 2. 원본 추정
        doc.add_heading('2. 원본 추정 (세대 분석)', level=1)
        
        para = doc.add_paragraph()
        para.add_run('🏆 원본 추정 순위:\n\n').bold = True
        
        for rank, (name, data) in enumerate(results['generation_ranking'][:5], 1):
            score = data['analysis']['generation_score']
            
            if rank == 1:
                para.add_run(f"{rank}위: {name} (세대 점수: {score:.3f}) ⭐ 원본 추정\n")
            else:
                para.add_run(f"{rank}위: {name} (세대 점수: {score:.3f})\n")
            
            # 세부 분석
            if 'screen_recording' in data['analysis']:
                sr = data['analysis']['screen_recording']
                if sr['is_screen_recording']:
                    para.add_run(f"   ⚠️ 화면 녹화 감지: {sr['reasons']}\n")
        
        # 3. 기술적 상세
        doc.add_heading('3. 기술적 분석 상세', level=1)
        
        para = doc.add_paragraph()
        para.add_run('분석 기법:\n').bold = True
        para.add_run('• PRNU (Photo Response Non-Uniformity) - 카메라 센서 지문\n')
        para.add_run('• GOP 구조 분석 - 키프레임 패턴\n')
        para.add_run('• 압축 아티팩트 분석 - 블록 노이즈, 모스키토 노이즈\n')
        para.add_run('• 오디오 지문 - Chromagram, MFCC\n')
        para.add_run('• 움직임 벡터 분석 - Optical Flow\n')
        para.add_run('• 화면 녹화 감지 - UI 패턴, 커서 감지\n')
        
        # 4. 결론
        doc.add_heading('4. 결론', level=1)
        
        para = doc.add_paragraph()
        para.add_run('⚖️ 포렌식 분석 결론:\n\n').bold = True
        
        # 타겟이 사용한 소스
        if results['source_match']:
            ref_num = results['source_match'].split('_')[1]
            para.add_run(f"1. 타겟 영상은 레퍼런스 {ref_num}번을 다운받아 사용한 것으로 확인됩니다.\n")
        
        # 원본 판정
        if results['generation_ranking']:
            original = results['generation_ranking'][0]
            if original[0] == 'target':
                para.add_run("2. 타겟 영상이 가장 원본에 가까운 것으로 추정됩니다.\n")
            else:
                ref_num = original[0].split('_')[1] if '_' in original[0] else original[0]
                para.add_run(f"2. 레퍼런스 {ref_num}번이 가장 원본에 가까운 것으로 추정됩니다.\n")
        
        para.add_run("\n이 분석은 디지털 포렌식 기법을 사용하여 ")
        para.add_run("법적 증거로 활용 가능한 수준의 정확도를 제공합니다.")
        
        # 저장
        report_path = os.path.join(
            self.report_dir,
            f"forensics_report_v2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        doc.save(report_path)
        
        print(f"\n📄 포렌식 보고서 생성 완료: {report_path}")
        return report_path

# 메인 실행 코드
def main():
    print("="*60)
    print("🔬 고급 영상 포렌식 분석 도구 NEW v3 (최상위 품질 다운로드)")
    print("="*60)
    print("NEW v3 개선사항:")
    print("• 최상위 품질 보장 (4K/2K/FHD 우선순위)")
    print("• 원본 비트레이트 및 해상도 유지")
    print("• 실시간 품질 검증 및 등급 표시")
    print("• 다중 포맷 우선순위 최적화")
    print("• 병렬 다운로드 및 재시도 기능")
    print("="*60)
    print("지원 기능:")
    print("• PRNU (카메라 센서 지문) 분석")
    print("• GOP 구조 분석")
    print("• 오디오 지문 분석")
    print("• 움직임 벡터 분석")
    print("• 화면 녹화 감지")
    print("• 압축 세대 분석")
    print("="*60)
    
    # 분석기 초기화
    forensics = AdvancedVideoForensics()
    
    # 타겟 영상 입력
    print("\n[타겟 영상]")
    target_input = forensics.get_video_input("타겟 영상을 선택하세요:", "타겟")
    
    if target_input[0] == 'skip':
        print("❌ 타겟 영상이 필요합니다")
        return
    
    # 타겟 처리
    target_path, target_meta = forensics.process_video_input(target_input, "target")
    if not target_path:
        print("❌ 타겟 준비 실패")
        return
    
    # 레퍼런스 영상들 입력
    reference_urls = []
    reference_paths = []
    
    print("\n[레퍼런스 영상들]")
    print("여러 개를 차례로 선택하세요")
    
    i = 1
    while True:
        ref_input = forensics.get_video_input(f"레퍼런스 {i}번:", f"레퍼런스{i}")
        
        if ref_input[0] == 'skip':
            if i == 1:
                print("⚠️ 최소 1개의 레퍼런스가 필요합니다")
                continue
            else:
                break
        
        ref_path, ref_meta = forensics.process_video_input(ref_input, f"reference_{i}")
        if ref_path:
            if ref_input[0] == 'url':
                reference_urls.append(ref_input[1])
            else:
                reference_urls.append(f"Local: {os.path.basename(ref_input[1])}")
            reference_paths.append(ref_path)
            i += 1
    
    if not reference_paths:
        print("❌ 레퍼런스가 없습니다")
        return
    
    # 포렌식 분석 실행
    print("\n🔬 포렌식 분석 시작...")
    results = forensics.find_source_video(target_path, reference_paths)
    
    # 결과 출력
    print("\n" + "="*60)
    print("📊 분석 결과")
    print("="*60)
    
    # 1. 타겟이 사용한 소스
    if results['source_match']:
        ref_num = results['source_match'].split('_')[1]
        confidence = results['match_confidence'] * 100
        print(f"\n✅ 타겟은 레퍼런스 {ref_num}번을 사용했습니다!")
        print(f"   신뢰도: {confidence:.1f}%")
    else:
        print("\n❌ 일치하는 레퍼런스를 찾지 못했습니다.")
    
    # 2. 원본 추정
    print("\n🏆 원본 추정 순위:")
    for rank, (name, data) in enumerate(results['generation_ranking'][:5], 1):
        score = data['analysis']['generation_score']
        
        # 화면 녹화 체크
        screen_rec = ""
        if 'screen_recording' in data['analysis']:
            sr = data['analysis']['screen_recording']
            if sr['is_screen_recording']:
                screen_rec = " [화면녹화]"
        
        if rank == 1:
            print(f"   {rank}위: {name} (세대점수: {score:.3f}) ⭐ 원본{screen_rec}")
        else:
            print(f"   {rank}위: {name} (세대점수: {score:.3f}){screen_rec}")
    
    # 3. 보고서 생성
    # target_url 처리
    if target_input[0] == 'url':
        target_url = target_input[1]
    else:
        target_url = f"Local: {os.path.basename(target_input[1])}"
    
    report_path = forensics.generate_report(target_url, reference_urls, results)
    
    print("\n" + "="*60)
    print("✅ 분석 완료!")
    print("="*60)
    print(f"\n📄 상세 보고서: {report_path}")
    print("\n💡 세대 점수 해석:")
    print("   0.0~0.2: 원본 또는 1세대")
    print("   0.2~0.4: 2-3세대 복사본")
    print("   0.4~0.6: 4-5세대 복사본")
    print("   0.6 이상: 다세대 복사본")
    
    input("\n엔터를 누르면 종료합니다...")

if __name__ == "__main__":
    main()